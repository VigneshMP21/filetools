# PDF Studio - Word Split Tool
#
# Supported modes:
#   page_breaks       -> Split DOCX using explicit/manual page breaks
#   sections          -> Split DOCX using Word section boundaries
#   selected_section  -> Extract one selected Word section
#   page_range        -> Extract visual page range using LibreOffice + pypdf
#
# The user selects the required mode at execution time.

import os
import shutil
import subprocess
import tempfile
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from pypdf import PdfReader, PdfWriter

from app.utils.file_validation import (
    validate_file_size,
    validate_word,
)


router = APIRouter(
    prefix="/api/word-tools/word-split",
    tags=["Word Tools"],
)


# ============================================================
# CONSTANTS
# ============================================================

WORD_NAMESPACE = (
    "http://schemas.openxmlformats.org/"
    "wordprocessingml/2006/main"
)

RELATIONSHIP_NAMESPACE = (
    "http://schemas.openxmlformats.org/"
    "officeDocument/2006/relationships"
)


# ============================================================
# COMMON HELPERS
# ============================================================

def get_libreoffice_path() -> str:
    """
    Find LibreOffice executable.

    Supports:
    - Windows development
    - Linux / Render
    """

    soffice = shutil.which("soffice")

    if soffice:
        return soffice

    libreoffice = shutil.which("libreoffice")

    if libreoffice:
        return libreoffice

    windows_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for path in windows_paths:
        if os.path.exists(path):
            return path

    raise HTTPException(
        status_code=500,
        detail="LibreOffice is not available on the server.",
    )


def remove_default_body_content(document: Document) -> None:
    """
    Remove the empty paragraph created automatically
    by python-docx while preserving sectPr.
    """

    body = document.element.body

    for child in list(body):

        if not child.tag.endswith("sectPr"):
            body.remove(child)


def create_docx_from_elements(elements) -> BytesIO:
    """
    Create a DOCX file from copied Word XML elements.
    """

    document = Document()

    remove_default_body_content(document)

    body = document.element.body

    sect_pr = body.sectPr

    insert_position = body.index(sect_pr)

    for element in elements:

        body.insert(
            insert_position,
            deepcopy(element),
        )

        insert_position += 1

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output


def create_zip_from_parts(parts, prefix="part") -> BytesIO:
    """
    Create ZIP containing multiple DOCX documents.
    """

    output = BytesIO()

    with ZipFile(
        output,
        mode="w",
        compression=ZIP_DEFLATED,
    ) as zip_file:

        for index, elements in enumerate(
            parts,
            start=1,
        ):

            document_output = create_docx_from_elements(
                elements
            )

            zip_file.writestr(
                f"{prefix}_{index}.docx",
                document_output.getvalue(),
            )

    output.seek(0)

    return output


# ============================================================
# MODE 1 - PAGE BREAKS
# ============================================================

def contains_page_break(element) -> bool:
    """
    Detect explicit/manual page breaks inside a paragraph.
    """

    if not element.tag.endswith("}p"):
        return False

    for break_element in element.iter(
        f"{{{WORD_NAMESPACE}}}br"
    ):

        break_type = break_element.get(
            f"{{{WORD_NAMESPACE}}}type"
        )

        if break_type == "page":
            return True

    return False


def split_by_page_breaks(document: Document):
    """
    Split document at explicit Ctrl + Enter page breaks.
    """

    parts = []
    current_part = []

    for element in document.element.body:

        if element.tag.endswith("sectPr"):
            continue

        current_part.append(element)

        if contains_page_break(element):

            parts.append(current_part)

            current_part = []

    if current_part:
        parts.append(current_part)

    if len(parts) <= 1:
        raise HTTPException(
            status_code=400,
            detail=(
                "No explicit page breaks were found "
                "in the Word document."
            ),
        )

    return parts


# ============================================================
# MODE 2 + MODE 3 - WORD SECTIONS
# ============================================================

def paragraph_has_section_break(element) -> bool:
    """
    Detect section properties stored inside paragraph pPr.

    In DOCX XML, a section boundary normally appears as:

        w:p
          w:pPr
            w:sectPr

    The final section's sectPr is normally stored directly
    under w:body.
    """

    if not element.tag.endswith("}p"):
        return False

    for child in element.iter():

        if child.tag == f"{{{WORD_NAMESPACE}}}sectPr":
            return True

    return False


def get_section_parts(document: Document):
    """
    Divide document body into logical Word sections.

    A section ends when a paragraph contains w:sectPr.
    The final body-level sectPr represents the final section
    configuration and is not treated as document content.
    """

    parts = []
    current_part = []

    for element in document.element.body:

        # Final body-level section properties.
        if element.tag.endswith("sectPr"):
            continue

        current_part.append(element)

        if paragraph_has_section_break(element):

            parts.append(current_part)

            current_part = []

    if current_part:
        parts.append(current_part)

    return parts


def split_by_sections(document: Document):
    """
    Mode 2:
    Split DOCX into all detected Word sections.
    """

    parts = get_section_parts(document)

    if len(parts) <= 1:
        raise HTTPException(
            status_code=400,
            detail=(
                "Only one Word section was detected. "
                "Add section breaks in Microsoft Word first."
            ),
        )

    return parts


def extract_selected_section(
    document: Document,
    section_number: int,
):
    """
    Mode 3:
    Extract one selected Word section.

    section_number is user-facing and starts at 1.
    """

    parts = get_section_parts(document)

    total_sections = len(parts)

    if section_number < 1:
        raise HTTPException(
            status_code=400,
            detail="Section number must start from 1.",
        )

    if section_number > total_sections:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Section {section_number} does not exist. "
                f"The document contains {total_sections} section(s)."
            ),
        )

    return parts[section_number - 1]


# ============================================================
# MODE 4 - VISUAL PAGE RANGE
# ============================================================

def convert_docx_to_pdf(
    docx_path: Path,
    output_directory: Path,
) -> Path:
    """
    Render DOCX as PDF using LibreOffice.

    This is required because DOCX itself does not contain
    reliable fixed visual page objects.
    """

    libreoffice = get_libreoffice_path()

    try:

        process = subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_directory),
                str(docx_path),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120,
            check=False,
        )

    except subprocess.TimeoutExpired:

        raise HTTPException(
            status_code=504,
            detail="Word page rendering timed out.",
        )

    pdf_path = (
        output_directory
        / f"{docx_path.stem}.pdf"
    )

    if (
        process.returncode != 0
        or not pdf_path.exists()
        or pdf_path.stat().st_size == 0
    ):

        raise HTTPException(
            status_code=500,
            detail="Unable to render the Word document.",
        )

    return pdf_path


def extract_pdf_page_range(
    pdf_path: Path,
    start_page: int,
    end_page: int,
) -> BytesIO:
    """
    Extract selected visual pages from rendered PDF.
    """

    reader = PdfReader(str(pdf_path))

    total_pages = len(reader.pages)

    if start_page < 1 or end_page < 1:

        raise HTTPException(
            status_code=400,
            detail="Page numbers must start from 1.",
        )

    if start_page > end_page:

        raise HTTPException(
            status_code=400,
            detail=(
                "Start page cannot be greater "
                "than end page."
            ),
        )

    if end_page > total_pages:

        raise HTTPException(
            status_code=400,
            detail=(
                f"Page range exceeds the document's "
                f"total of {total_pages} pages."
            ),
        )

    writer = PdfWriter()

    for page_index in range(
        start_page - 1,
        end_page,
    ):

        writer.add_page(
            reader.pages[page_index]
        )

    output = BytesIO()

    writer.write(output)

    output.seek(0)

    return output


# ============================================================
# MAIN API
# ============================================================

@router.post("/")
async def word_split(
    file: UploadFile = File(...),

    mode: str = Form(...),

    section_number: int | None = Form(None),

    start_page: int | None = Form(None),

    end_page: int | None = Form(None),
):
    """
    Word Split API.

    Available modes:

    page_breaks
        Split document using explicit page breaks.

    sections
        Split document using Word section breaks.

    selected_section
        Extract one selected Word section.

    page_range
        Render Word document using LibreOffice and
        extract selected visual pages as PDF.
    """

    await validate_file_size(file)
    await validate_word(file)

    mode = mode.strip().lower()

    valid_modes = {
        "page_breaks",
        "sections",
        "selected_section",
        "page_range",
    }

    if mode not in valid_modes:

        raise HTTPException(
            status_code=400,
            detail=(
                "Invalid split mode. "
                "Use page_breaks, sections, "
                "selected_section, or page_range."
            ),
        )

    temp_directory = None

    try:

        await file.seek(0)

        content = await file.read()

        # ====================================================
        # MODE 1
        # ====================================================

        if mode == "page_breaks":

            document = Document(
                BytesIO(content)
            )

            parts = split_by_page_breaks(
                document
            )

            zip_output = create_zip_from_parts(
                parts,
                prefix="page_break_part",
            )

            return StreamingResponse(
                zip_output,
                media_type="application/zip",
                headers={
                    "Content-Disposition":
                        'attachment; filename="word-page-breaks.zip"'
                },
            )

        # ====================================================
        # MODE 2
        # ====================================================

        if mode == "sections":

            document = Document(
                BytesIO(content)
            )

            parts = split_by_sections(
                document
            )

            zip_output = create_zip_from_parts(
                parts,
                prefix="section",
            )

            return StreamingResponse(
                zip_output,
                media_type="application/zip",
                headers={
                    "Content-Disposition":
                        'attachment; filename="word-sections.zip"'
                },
            )

        # ====================================================
        # MODE 3
        # ====================================================

        if mode == "selected_section":

            if section_number is None:

                raise HTTPException(
                    status_code=400,
                    detail=(
                        "section_number is required "
                        "for selected_section mode."
                    ),
                )

            document = Document(
                BytesIO(content)
            )

            elements = extract_selected_section(
                document,
                section_number,
            )

            output = create_docx_from_elements(
                elements
            )

            return StreamingResponse(
                output,
                media_type=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),
                headers={
                    "Content-Disposition":
                        (
                            "attachment; "
                            f'filename="section_{section_number}.docx"'
                        )
                },
            )

        # ====================================================
        # MODE 4
        # ====================================================

        if mode == "page_range":

            if start_page is None or end_page is None:

                raise HTTPException(
                    status_code=400,
                    detail=(
                        "start_page and end_page are required "
                        "for page_range mode."
                    ),
                )

            temp_directory = tempfile.mkdtemp(
                prefix="pdfstudio_word_split_"
            )

            temp_path = Path(
                temp_directory
            )

            input_path = (
                temp_path / "input.docx"
            )

            with open(
                input_path,
                "wb",
            ) as output_file:

                output_file.write(content)

            rendered_pdf = convert_docx_to_pdf(
                input_path,
                temp_path,
            )

            pdf_output = extract_pdf_page_range(
                rendered_pdf,
                start_page,
                end_page,
            )

            return StreamingResponse(
                pdf_output,
                media_type="application/pdf",
                headers={
                    "Content-Disposition":
                        (
                            "attachment; "
                            f'filename="pages_{start_page}-{end_page}.pdf"'
                        )
                },
            )

        raise HTTPException(
            status_code=400,
            detail="Unsupported Word split mode.",
        )

    except HTTPException:
        raise

    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Unable to split the Word document.",
        )

    finally:

        await file.close()

        if (
            temp_directory
            and os.path.exists(temp_directory)
        ):

            shutil.rmtree(
                temp_directory,
                ignore_errors=True,
            )