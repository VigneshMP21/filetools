#!/usr/bin/env python3
"""Convert a PDF into a Word document (.docx) by extracting text per page."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Iterable, List, Optional, Sequence


def _load_dependencies() -> tuple[object, object]:
    try:
        import pypdf  # type: ignore
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise RuntimeError(
            "Missing dependency: install 'pypdf' to extract text from PDF files."
        ) from exc

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise RuntimeError(
            "Missing dependency: install 'python-docx' to create Word documents."
        ) from exc

    return pypdf, Document


def _clean_text(raw_text: Optional[str]) -> List[str]:
    if not raw_text:
        return []

    cleaned: List[str] = []
    for line in raw_text.splitlines():
        stripped = line.strip()
        if stripped:
            cleaned.append(stripped)
    return cleaned


def convert_pdf_to_word(input_pdf: str | Path, output_docx: str | Path) -> Path:
    input_path = Path(input_pdf)
    output_path = Path(output_docx)

    if not input_path.exists():
        raise FileNotFoundError(f"Input PDF not found: {input_path}")

    pypdf, Document = _load_dependencies()
    reader = pypdf.PdfReader(str(input_path))

    if not reader.pages:
        raise ValueError("The PDF does not contain any pages.")

    document = Document()
    document.add_heading("Converted PDF", level=1)
    document.add_paragraph(f"Source file: {input_path.name}")

    for index, page in enumerate(reader.pages, start=1):
        document.add_heading(f"Page {index}", level=2)
        text_blocks = _clean_text(page.extract_text())

        if not text_blocks:
            document.add_paragraph("[No selectable text was found on this page.]")
            continue

        for paragraph in text_blocks:
            document.add_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Convert a PDF to a Word document")
    parser.add_argument("input_pdf", help="Path to the source PDF file")
    parser.add_argument("output_docx", help="Path to the generated .docx file")
    args = parser.parse_args(list(argv) if argv is not None else None)

    try:
        output_path = convert_pdf_to_word(args.input_pdf, args.output_docx)
    except Exception as exc:  # pragma: no cover - exercised via CLI usage
        print(f"Conversion failed: {exc}", file=sys.stderr)
        return 1

    print(f"Created Word file: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
